from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_DIR = Path(
    r"C:\Users\ondra\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents"
)
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


ROOT = Path(r"C:\Users\ondra\Desktop\KSP Modding\KickLifeSupportFork")
OUTPUT = ROOT / "KICK_Quick_Reference.docx"
PAGE_WIDTH_DXA = 9360
TITLE_COLOR = "1F4D78"
HEADING_COLOR = "2E74B5"
HEADER_FILL = "E8EEF5"


def set_font(run, size, *, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, before, after, color in [
        ("Heading 1", 15, 14, 6, HEADING_COLOR),
        ("Heading 2", 12.5, 10, 4, TITLE_COLOR),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.05


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
    return paragraph


def set_cell_text(cell, text, *, bold=False, size=9.5, color=None):
    paragraph = clear_paragraph(cell.paragraphs[0])
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    lines = str(text).split("\n")
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_font(run, size, bold=bold, color=color)
        if idx < len(lines) - 1:
            run.add_break()
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("KICK Life Support Quick Reference")
    set_font(run, 19, bold=True, color=TITLE_COLOR)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    subrun = subtitle.add_run(
        "Current systems snapshot for testing and patching - generated "
        + datetime.now().strftime("%Y-%m-%d %H:%M")
    )
    set_font(subrun, 10, color="555555")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    apply_table_geometry(table, widths, table_width_dxa=PAGE_WIDTH_DXA)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, size=9.5, color=TITLE_COLOR)
        shade_cell(cell, HEADER_FILL)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    apply_table_geometry(table, widths, table_width_dxa=PAGE_WIDTH_DXA)
    return table


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_font(run, 10)


def build_doc():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_paragraph(
        "This sheet is the short operational version: what each cabin/control option does, "
        "what the important toggles mean, and which limits currently matter in flight."
    )

    doc.add_heading("Core Rules", level=1)
    add_bullets(
        doc,
        [
            "Ambient breathing only counts when the body atmosphere contains oxygen and outside pressure is at least 50.7 kPa.",
            "Empty retained cabins with CO2 still present now scrub at a minimum one-seat rate.",
            "Atmospheric Control 'Use Max Capacity' forces full installed scrubber capacity instead of occupancy-scaled capacity.",
            "Cabin-part conductance is 0.1 for Unpressurized, 0.01 for Open-Loop Ventilation, and 0.001 for sealed cabins.",
            "Air Cooling starts at 1.0 kPa and reaches full output at 50.7 kPa.",
            "Water Evaporator only works below 0.6 kPa; WasteWater provides 80% cooling and leaves 20% Waste.",
        ],
    )

    doc.add_heading("Atmospheric Control", level=1)
    atcon_rows = [
        [
            "Unpressurized Cabin",
            "26 kPa",
            "Yes",
            "No",
            "None",
            "0",
            "None",
            "No shared CO2 cabin; immediate ambient dependence.",
        ],
        [
            "Open-Loop Ventilation",
            "2 kPa",
            "Yes",
            "Yes",
            "0.005 L/s per seat",
            "0.005 EC/s per occupied seat",
            "Oxygen up to 11 units per unit CO2 removed outside ambient breathing",
            "Powered venting in thin air or vacuum; can also use ambient air when available.",
        ],
        [
            "Pressurized Cabin",
            "Vacuum rated",
            "No",
            "Yes",
            "None",
            "0",
            "None",
            "Sealed volume only; needs another cabin to scrub CO2.",
        ],
        [
            "LiOH Scrubber",
            "Vacuum rated",
            "No",
            "Yes",
            "0.005 L/s per seat",
            "0.050 EC/s per occupied seat",
            "LithiumHydroxide",
            "Creates electrical heat and reaction heat; produces Waste.",
        ],
        [
            "Zeolite Molecular Sieve",
            "Vacuum rated",
            "No",
            "Yes",
            "0.005 L/s per seat",
            "0.200 EC/s per occupied seat",
            "None",
            "Regenerative scrubber; dumps collected CO2 to the vessel resource pool.",
        ],
        [
            "Solid Amine Swingbed",
            "Vacuum rated",
            "No",
            "Yes",
            "0.005 L/s per seat",
            "0.100 EC/s per occupied seat",
            "None",
            "Lighter regenerative option with lower EC demand.",
        ],
    ]
    add_table(
        doc,
        [
            "System",
            "Pressure Rating",
            "Ambient Breathes",
            "Retains CO2",
            "CO2 Removal",
            "EC",
            "Consumables",
            "Notes",
        ],
        atcon_rows,
        [1500, 980, 760, 760, 1220, 1100, 1280, 1760],
    )

    doc.add_heading("Environmental Control", level=1)
    encon_rows = [
        [
            "Passive",
            "None",
            "0",
            "None",
            "Relies on cabin-part exchange only",
        ],
        [
            "Air Cooling",
            "1.00 kW per seat",
            "0.200 EC/s per seat",
            "None",
            "Atmosphere sink: ramps from 1.0 to 50.7 kPa",
        ],
        [
            "Water Evaporator",
            "1.00 kW per seat",
            "0.010 EC/s per seat",
            "Water or WasteWater up to 0.000435 L/s per seat",
            "Vacuum/near-vacuum only: below 0.6 kPa",
        ],
        [
            "Pumped Coolant Loop",
            "1.00 kW per seat",
            "0.100 EC/s per seat",
            "None",
            "Needs external SystemHeat radiators to reject loop heat",
        ],
    ]
    add_table(
        doc,
        ["System", "Cooling", "Cooling EC", "Consumables", "Notes"],
        encon_rows,
        [1700, 1400, 1200, 2100, 2960],
    )

    doc.add_heading("Important PAW Toggles", level=1)
    ui_rows = [
        ["Atmospheric Control", "CO2 Removal", "Master on/off for any powered atmospheric control system."],
        ["Atmospheric Control", "Use Max Capacity", "Uses full installed scrubber capacity instead of occupancy-scaled capacity."],
        ["Environmental Control", "Master Switch", "Enables the thermal-control package and its base EC draw."],
        ["Environmental Control", "Heater", "Allows active cabin heating up to the heater limit."],
        ["Environmental Control", "Cooling", "Runs the installed cooling system when conditions allow it."],
        ["Environmental Control", "Evaporator Feed", "Water / WasteWater priority selection for Water Evaporator only."],
        ["Environmental Control", "EC Limit", "Caps total active thermal-control EC use."],
        ["Environmental Control", "Water Limit", "Caps Water/WasteWater consumption for evaporative cooling."],
    ]
    add_table(
        doc,
        ["Section", "Control", "What It Does"],
        ui_rows,
        [1500, 1700, 6160],
    )

    doc.add_heading("Safety Numbers", level=1)
    safety_rows = [
        ["CO2 warning", "3%"],
        ["CO2 fatal", "10%"],
        ["Safe cabin temperature", "5 C to 45 C"],
        ["Warning cabin temperature", "10 C to 40 C"],
        ["Oxygen grace period", "120 s"],
        ["Pressure-failure grace period", "60 s"],
    ]
    add_table(doc, ["Threshold", "Current Value"], safety_rows, [4200, 5160])

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(str(OUTPUT))
