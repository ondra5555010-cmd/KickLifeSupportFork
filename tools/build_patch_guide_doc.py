from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_DIR = Path(
    r"C:\Users\ondra\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents"
)
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


OUTPUT = Path(
    r"C:\Users\ondra\Desktop\KSP Modding\KickLifeSupportFork\KICK_Life_Support_Config_and_Patching_Guide.docx"
)

PAGE_WIDTH_DXA = 9360
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def clear_cell(cell):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        p._element.remove(run._element)
    return p


def set_paragraph_text(paragraph, text, *, bold=False, font_size=10.5, color=None):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_multiline_cell(cell, text, *, bold=False, font_size=10.5, color=None):
    lines = text.split("\n")
    p = clear_cell(cell)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if idx < len(lines) - 1:
            run.add_break()
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(subtitle)
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r2.font.size = Pt(11)
    r2.italic = True
    r2.font.color.rgb = RGBColor.from_string("555555")


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run("Note: ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(10.5)
    r1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r2 = p.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    apply_table_geometry(table, widths, table_width_dxa=PAGE_WIDTH_DXA)
    for idx, header in enumerate(headers):
        add_multiline_cell(table.rows[0].cells[idx], header, bold=True, font_size=10.5, color=DARK_BLUE)
        set_cell_shading(table.rows[0].cells[idx], HEADER_FILL)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            add_multiline_cell(cells[idx], value)

    apply_table_geometry(table, widths, table_width_dxa=PAGE_WIDTH_DXA)
    return table


def add_system_table(doc, title, summary, rows):
    doc.add_heading(title, level=3)
    if summary:
        doc.add_paragraph(summary)
    add_table(doc, ["Field", "Value", "Note"], rows, [2100, 2400, 4860])


def add_settings_section(doc):
    doc.add_heading("CONFIG FILES", level=1)
    doc.add_paragraph(
        "This section documents the global config node in Settings.cfg and the most useful patchable fields on KickAvionicsModule and KickLifeSupportModule. "
        "Runtime UI/status fields are omitted unless they are useful to seed defaults from part configs or B9 subtype DATA blocks."
    )

    doc.add_heading("Settings.cfg", level=2)
    doc.add_paragraph(
        "KickLifeSupportConfig reads the first KICKLS_SETTINGS node in GameData and exposes values through GetDouble/GetFloat. "
        "Changing these values affects every part that uses the shared life-support and thermal code."
    )

    sections = [
        (
            "Consumption and production",
            [
                ("OXYGEN_RATE", "0.005", "Base oxygen use per kerbal per second."),
                ("WATER_RATE", "0.00003", "Base water use per kerbal per second."),
                ("FOOD_RATE", "0.00002", "Base food use per kerbal per second."),
                ("CO2_RATE", "0.0041", "CO2 generated per kerbal per second."),
                ("SCRUBBER_RATE", "0.005", "Nominal CO2 removal capacity per occupied seat."),
                ("WASTEWATER_RATE", "0.000028", "WasteWater generated per kerbal per second."),
                ("WASTE_RATE", "0.000018", "Solid waste generated per kerbal per second."),
                ("LITHIUMHYDROXIDE_RATE", "0.000015", "LiOH consumed per unit of scrubber work."),
            ],
        ),
        (
            "Grace periods and ambient rules",
            [
                ("GRACE_OXYGEN", "120 s", "Time before crew dies after losing breathable air."),
                ("GRACE_UNPRESSURIZED", "60 s", "Time before pressure-limited cabins fail when exposed below their limit."),
                ("GRACE_WATER", "129600 s", "Water deprivation grace period."),
                ("GRACE_FOOD", "604800 s", "Food deprivation grace period."),
                ("GRACE_CLIMATE", "3600 s", "Temperature-control grace period used by the life-support scenario."),
                ("GRACE_TEMP", "300 s", "Direct lethal cabin-temperature grace period."),
                ("AMBIENT_PRESSURE_MINIMUM", "50.7 kPa", "Minimum oxygen-bearing outside pressure that counts as a safe ambient atmosphere."),
            ],
        ),
        (
            "Thermal model and heat rejection",
            [
                ("KERBAL_HEAT", "0.03 kW", "Heat produced by each kerbal into the cabin."),
                ("LIOH_REACTION_HEAT_PER_UNIT", "4.0", "Reaction heat added by LiOH scrubbing per unit of scrubbed CO2."),
                ("AIR_DENSITY", "0.001225", "Assumed cabin-air density used to derive cabin air mass from airVolumePerSeat."),
                ("AIR_SPECIFIC_HEAT", "1005", "Specific heat used for the cabin-air component of thermal mass."),
                ("GENERIC_CABIN_SPECIFIC_HEAT", "1000", "Specific heat used for the part-mass fraction added to cabin thermal mass."),
                ("WASTEWATER_COOLING_FACTOR", "0.8", "WasteWater cooling efficiency relative to clean water."),
                ("WASTEWATER_RESIDUE_MASS_FRACTION", "0.2", "Fraction of WasteWater mass returned as Waste after evaporation."),
                ("EVAPORATION_ENERGY_KJ_PER_UNIT", "2300", "Cooling energy extracted per unit of evaporated water."),
                ("WATER_EVAPORATOR_PRESSURE_LIMIT", "0.6 kPa", "Maximum outside pressure at which the evaporator may run."),
                ("AIR_COOLING_MIN_PRESSURE", "1.0 kPa", "Pressure where Air Cooling begins to produce any output."),
                ("AIR_COOLING_FULL_PRESSURE", "50.7 kPa", "Pressure where Air Cooling reaches full rated output."),
                ("HIGH_WARP_STABILIZATION_THRESHOLD", "100x", "Warp threshold above which the simplified high-warp temperature stabilizer is allowed to help."),
            ],
        ),
        (
            "Warnings and PAW slider ranges",
            [
                ("CO2_WARNING_LEVEL", "0.03", "CO2 fraction where the warning state begins."),
                ("CO2_FATAL_LEVEL", "0.10", "CO2 fraction where the fatal state begins."),
                ("MIN_SAFE_CABIN_TEMP", "5 C", "Cabin temperature below which the situation becomes unsafe."),
                ("MAX_SAFE_CABIN_TEMP", "45 C", "Cabin temperature above which the situation becomes unsafe."),
                ("MIN_WARNING_CABIN_TEMP", "10 C", "Low-temperature warning threshold."),
                ("MAX_WARNING_CABIN_TEMP", "40 C", "High-temperature warning threshold."),
                ("THERMOSTAT_MIN / MAX / DEFAULT / STEP", "10 / 30 / 22 / 0.5", "Shared thermostat slider limits and default value."),
                ("ENCON_EC_LIMIT_STEP", "0.005", "Step size for the EnCon EC limit slider."),
                ("EVAPORATOR_WATER_LIMIT_STEP", "0.00005", "Step size for the evaporator water-limit slider."),
            ],
        ),
    ]

    for heading, rows in sections:
        doc.add_heading(heading, level=3)
        add_table(doc, ["Field", "Current Default", "What It Does"], rows, [2200, 1400, 5760])

    doc.add_heading("KickAvionicsModule", level=2)
    doc.add_paragraph(
        "KickAvionicsModule is the simple electrical-control module used for command, SAS, and RCS support. "
        "In stock Kick Life Support, 30_CapsuleSystems.cfg seeds its default rates by CrewCapacity."
    )
    add_table(
        doc,
        ["Field", "Meaning", "Typical Source / Notes"],
        [
            ("avionicsEnabled", "Default state of the avionics master switch.", "Can be seeded in the part config or a patch."),
            ("avionicsECRate", "Base EC drain for basic command operation.", "30_CapsuleSystems.cfg sets 0.02 EC/s for 1-3 seats and 0.2 EC/s for >3 seats."),
            ("avionicsHeat", "Heat produced by the base avionics load.", "Normally patched to 0.02 kW in the supplied defaults."),
            ("sasECRate", "Extra EC drain while SAS is active.", "Default patch uses 0.01 EC/s for 1-3 seats and 0.1 EC/s for >3 seats."),
            ("sasHeat", "Heat produced by SAS operation.", "Default patch uses 0.01 kW."),
            ("rcsECRate", "Extra EC drain while RCS is active.", "Default patch uses 0.01 EC/s for 1-3 seats and 0.1 EC/s for >3 seats."),
            ("rcsHeat", "Heat produced by RCS control operation.", "Default patch uses 0.01 kW."),
        ],
        [1900, 3300, 4160],
    )
    add_note(
        doc,
        "dbsAvionicsECRate and the visible PAW display fields are runtime outputs. They are useful for UI/DBS integration, but they are not the values you normally patch."
    )

    doc.add_heading("KickLifeSupportModule", level=2)
    doc.add_paragraph(
        "KickLifeSupportModule is split between KickLifeSupportModule.cs (atmospheric control, cabin metrics, life support, PAW state) "
        "and KickLifeSupportModule.Thermal.cs (thermal model and Environmental Control / EnCon behavior). "
        "Most part patches only need a small subset of the fields below."
    )

    doc.add_heading("Atmospheric-control and cabin fields", level=3)
    add_table(
        doc,
        ["Field", "Meaning", "Notes"],
        [
            ("lifeSupportEnabled", "Master enable for life-support logic on the part.", "Usually true on crewed parts."),
            ("atmosphereControlMode", "Atmospheric-control mode enum.", "0=None, 1=OpenLoop, 2=LiOH, 3=Regenerative."),
            ("atmosphereControlSystemName", "Display name shown in the Atmospheric Control PAW.", "Often set by B9 subtype DATA."),
            ("atmosphericControlECRate", "Rated EC draw per occupied-seat equivalent for the current atmospheric-control system.", "Scaled in runtime by occupancy and system rules."),
            ("atmosphericControlHeatPerEC", "Electrical heat added per EC spent by atmospheric control.", "Current B9 subtypes use 0 for passive/pressurized and 1 for powered systems."),
            ("oxygenWastePerCO2Removed", "Extra oxygen consumed per unit of CO2 removed outside safe ambient conditions.", "Open-loop uses 11; sealed systems use 0."),
            ("useMaxCapacity", "Default state of the Atmospheric Control 'Use Max Capacity' toggle.", "False by default; when enabled the scrubber uses full installed capacity instead of occupancy-scaled capacity."),
            ("pressureMinimumKPa", "Minimum outside pressure the cabin can tolerate before depressurization grace begins.", "0 means vacuum-rated / no pressure floor."),
            ("canUseAmbient", "Whether breathable outside air can satisfy crew breathing when the atmosphere is safe.", "True on unpressurized and open-loop cabins."),
            ("retainsCO2", "Whether the part participates in the vessel-wide sealed CO2 pool.", "False forces effective atmosphere mode to None."),
            ("cabinMassFraction", "Fraction of dry part mass folded into cabin thermal mass.", "Current default is 0.05."),
            ("airVolumePerSeat", "Cabin air volume assigned per crew seat.", "Current default is 2000 volume units per seat."),
            ("cabinPartConductance", "Passive thermal coupling between cabin air and stock part temperature.", "Current AtCon switch uses 0.1 / 0.01 / 0.001 depending on subtype."),
        ],
        [1850, 3300, 4210],
    )

    doc.add_heading("Environmental-control and thermal behavior fields", level=3)
    add_table(
        doc,
        ["Field", "Meaning", "Notes"],
        [
            ("systemECRate", "Base EnCon package EC draw.", "Applied whenever EnCon master switch is on."),
            ("systemHeat", "Base EnCon package heat added into the cabin.", "Current default patch uses 0.003 kW."),
            ("enconMaxECRate", "Maximum equipment EC budget available to heater/cooler hardware on the part.", "Current default patch sets this to CrewCapacity."),
            ("heaterMaxECRatePerSeat", "Heater EC budget per seat.", "Current default is 1.0 EC/s per seat."),
            ("heaterHeatPerEC", "Heater output in kW per 1 EC/s consumed.", "Current default is 1.0."),
            ("coolingSystemMode", "Cooling system enum.", "0=None, 1=Integrated, 2=Coolant Loop."),
            ("coolingSystemName", "Display name for the current cooler.", "Usually set by B9 subtype DATA."),
            ("coolingMaxECRatePerSeat", "Rated cooling EC budget per seat.", "Examples: 0.2 Air Cooling, 0.01 Water Evaporator, 0.1 Pumped Coolant Loop."),
            ("coolingHeatPerEC", "Cooling capacity in kW per 1 EC/s.", "Examples: 5 Air Cooling, 100 Water Evaporator, 10 Pumped Coolant Loop."),
            ("coolingMaxWaterRatePerSeat", "Maximum water or wastewater consumption per seat for consumable integrated cooling.", "Zero on Air Cooling and Pumped Coolant Loop."),
            ("coolingPressureMinimumKPa", "Pressure floor used by the active cooler.", "Air Cooling starts above this; Water Evaporator shuts off at and above this."),
            ("coolingPressureFullKPa", "Pressure where an atmospheric integrated cooler reaches full rated output.", "Unused by Water Evaporator and Pumped Coolant Loop."),
            ("heatPumpLoopNominalTemp", "SystemHeat loop temperature where the heat pump still runs at full effectiveness.", "Default 320 K."),
            ("heatPumpLoopMaxTemp", "SystemHeat loop temperature where the heat pump falls to zero effectiveness.", "Default 370 K."),
            ("heatPumpSystemHeatModuleID", "Target SystemHeat moduleID used as the part's coolant-loop entry point.", "90_Compat_SystemHeat.cfg sets this to kickECS."),
        ],
        [1850, 3300, 4210],
    )

    doc.add_heading("Persisted PAW defaults worth seeding from configs", level=3)
    add_table(
        doc,
        ["Field", "What It Seeds", "Notes"],
        [
            ("scrubberEnabled", "Default Atmospheric Control master-switch state.", "Usually left on."),
            ("useMaxCapacity", "Default Atmospheric Control capacity-mode toggle.", "False by default."),
            ("climateControlEnabled", "Default EnCon master-switch state.", "Usually left on."),
            ("thermostatTemp", "Default EnCon thermostat value.", "30_CapsuleSystems.cfg seeds this from THERMOSTAT_DEFAULT."),
            ("enconECLimit", "Default EnCon equipment EC slider value.", "Current patch seeds this to CrewCapacity."),
            ("heaterEnabled", "Default heater toggle state.", "Persisted and patchable."),
            ("airCoolingEnabled", "Default atmospheric integrated-cooling toggle state.", "Only visible on dry Integrated subtypes such as Air Cooling."),
            ("evaporatorEnabled", "Default consumable integrated-cooling toggle state.", "Only visible on wet Integrated subtypes such as Water Evaporator."),
            ("heatPumpEnabled", "Default coolant-loop toggle state.", "Only visible on Pumped Coolant Loop subtype."),
            ("evaporatorFeedMode", "Default evaporator feed priority.", "Waste First, Waste Only, Fresh First, or Fresh Only."),
            ("waterEvaporatorWaterLimit", "Default evaporator water-limit slider value.", "Base EnCon patch seeds it from CrewCapacity * 0.000435."),
        ],
        [1850, 3300, 4210],
    )


def add_b9_section(doc):
    doc.add_heading("B9 SYSTEM SWITCH", level=1)
    doc.add_paragraph(
        "Kick Life Support uses two B9PartSwitch modules. Atmospheric Control is defined in 20_AtmosphericControl_B9.cfg and requires only B9PartSwitch. "
        "Environmental Control - Systems is defined in 21_ThermalControl_B9.cfg and currently requires both SystemHeat and B9PartSwitch."
    )

    doc.add_heading("Atmospheric Control (AtCon)", level=2)
    add_note(
        doc,
        "Before any subtype DATA is applied, the AtCon patch seeds module defaults such as cabinMassFraction=0.05, airVolumePerSeat=2000, cabinPartConductance=0.001, retainsCO2=true, and canUseAmbient=false."
    )

    add_system_table(
        doc,
        "Unpressurized Cabin",
        "Always available.",
        [
            ("atmosphereControlMode", "0", ""),
            ("atmosphereControlSystemName", "Unpressurized Cabin", ""),
            ("pressureMinimumKPa", "26", ""),
            ("cabinPartConductance", "0.1", ""),
            ("canUseAmbient", "true", ""),
            ("retainsCO2", "false", ""),
            ("addedMass", "-10% of part mass", ""),
            ("addedCost", "-20% of part cost", ""),
        ],
    )
    add_system_table(
        doc,
        "Open-Loop Ventilation",
        "upgradeRequired = kick-life-support-upgrade-OpenLoopVenting",
        [
            ("atmosphereControlMode", "1", ""),
            ("atmosphereControlSystemName", "Open-Loop Ventilation", ""),
            ("atmosphericControlECRate", "0.005", "Per occupied seat."),
            ("atmosphericControlHeatPerEC", "1.0", ""),
            ("oxygenWastePerCO2Removed", "11", ""),
            ("useMaxCapacity", "false", "Default toggle state."),
            ("pressureMinimumKPa", "2", ""),
            ("cabinPartConductance", "0.01", ""),
            ("canUseAmbient", "true", ""),
            ("retainsCO2", "true", ""),
            ("RESOURCE Oxygen", "unitsPerVolume = 200", ""),
            ("defaultSubtypePriority", "1", ""),
            ("addedMass", "-5% of part mass", ""),
            ("addedCost", "-7.5% of part cost", ""),
        ],
    )
    add_system_table(
        doc,
        "Pressurized Cabin",
        "upgradeRequired = kick-life-support-upgrade-PressurizedCabin",
        [
            ("atmosphereControlMode", "0", ""),
            ("atmosphereControlSystemName", "Pressurized Cabin", ""),
            ("canUseAmbient", "false", ""),
            ("retainsCO2", "true", ""),
            ("RESOURCE Oxygen", "unitsPerVolume = 200", ""),
            ("addedMass", "-3% of part mass", ""),
            ("addedCost", "-5% of part cost", ""),
        ],
    )
    add_system_table(
        doc,
        "LiOH Scrubber",
        "upgradeRequired = kick-life-support-upgrade-LiOH",
        [
            ("atmosphereControlMode", "2", ""),
            ("atmosphereControlSystemName", "LiOH Scrubber", ""),
            ("atmosphericControlECRate", "0.05", "Per occupied seat."),
            ("atmosphericControlHeatPerEC", "1.0", ""),
            ("RESOURCE Oxygen", "unitsPerVolume = 200", ""),
            ("RESOURCE LithiumHydroxide", "unitsPerVolume = 1", ""),
        ],
    )
    add_system_table(
        doc,
        "Zeolite Molecular Sieve",
        "upgradeRequired = kick-life-support-upgrade-Zeolite",
        [
            ("atmosphereControlMode", "3", ""),
            ("atmosphereControlSystemName", "Zeolite Molecular Sieve", ""),
            ("atmosphericControlECRate", "0.20", "Per occupied seat."),
            ("atmosphericControlHeatPerEC", "1.0", ""),
            ("RESOURCE Oxygen", "unitsPerVolume = 200", ""),
            ("addedMass", "0.075 + 0.025 per seat", ""),
            ("addedCost", "750 + 250 per seat", ""),
        ],
    )
    add_system_table(
        doc,
        "Solid Amine Swingbed",
        "upgradeRequired = kick-life-support-upgrade-SolidAmine",
        [
            ("atmosphereControlMode", "3", ""),
            ("atmosphereControlSystemName", "Solid Amine Swingbed", ""),
            ("atmosphericControlECRate", "0.10", "Per occupied seat."),
            ("atmosphericControlHeatPerEC", "1.0", ""),
            ("RESOURCE Oxygen", "unitsPerVolume = 200", ""),
            ("addedMass", "0.05 + 0.015 per seat", ""),
            ("addedCost", "1250 + 750 per seat", ""),
        ],
    )

    doc.add_heading("Environmental Control - Systems (EnCon)", level=2)
    add_note(
        doc,
        "The EnCon B9 patch only chooses the active cooling hardware. Separate defaults still come from 30_CapsuleSystems.cfg: systemECRate=0.003, systemHeat=0.003, enconMaxECRate=CrewCapacity, enconECLimit=CrewCapacity, heaterMaxECRatePerSeat=1.0, and heaterHeatPerEC=1.0."
    )
    add_system_table(
        doc,
        "Passive",
        "Always available.",
        [
            ("coolingSystemMode", "0", ""),
            ("coolingSystemName", "Passive", ""),
            ("addedCost", "-5% of part cost", ""),
        ],
    )
    add_system_table(
        doc,
        "Air Cooling",
        "upgradeRequired = kick-life-support-upgrade-AirCooling",
        [
            ("coolingSystemMode", "1", ""),
            ("coolingSystemName", "Air Cooling", ""),
            ("coolingMaxECRatePerSeat", "0.2", ""),
            ("coolingHeatPerEC", "5", ""),
            ("coolingPressureMinimumKPa", "AIR_COOLING_MIN_PRESSURE", "Currently 1.0 kPa."),
            ("coolingPressureFullKPa", "AIR_COOLING_FULL_PRESSURE", "Currently 50.7 kPa."),
        ],
    )
    add_system_table(
        doc,
        "Water Evaporator",
        "upgradeRequired = kick-life-support-upgrade-WaterEvaporator",
        [
            ("coolingSystemMode", "1", ""),
            ("coolingSystemName", "Water Evaporator", ""),
            ("coolingMaxECRatePerSeat", "0.01", ""),
            ("coolingHeatPerEC", "100", ""),
            ("coolingMaxWaterRatePerSeat", "0.000435", ""),
            ("coolingPressureMinimumKPa", "WATER_EVAPORATOR_PRESSURE_LIMIT", "Currently 0.6 kPa."),
            ("waterEvaporatorWaterLimit", "CrewCapacity * 0.000435", "Base default."),
        ],
    )
    add_system_table(
        doc,
        "Pumped Coolant Loop",
        "upgradeRequired = kick-life-support-upgrade-CoolantLoop",
        [
            ("coolingSystemMode", "2", ""),
            ("coolingSystemName", "Pumped Coolant Loop", ""),
            ("coolingMaxECRatePerSeat", "0.1", ""),
            ("coolingHeatPerEC", "10", ""),
            ("addedCost", "250 + 250 per seat", ""),
            ("ModuleSystemHeat.moduleID", "kickECS", "From 90_Compat_SystemHeat.cfg."),
            ("ModuleSystemHeat.volume", "CrewCapacity * 0.001", "Only when SystemHeat is present."),
        ],
    )


def add_tech_tree_section(doc):
    doc.add_heading("TECH TREE ASSIGNMENT", level=1)
    doc.add_paragraph(
        "Tech-gating is defined through PARTUPGRADE nodes in 00_Upgrades.cfg. "
        "A B9 subtype becomes tech-gated by adding upgradeRequired = <upgrade name>. "
        "After the tech node is researched, the player still pays the PARTUPGRADE entryCost to unlock the subtype family."
    )

    add_table(
        doc,
        ["Upgrade name", "Shown title", "Tech node", "Entry cost", "Used by subtype(s)"],
        [
            ("kick-life-support-upgrade-OpenLoopVenting", "Open-Loop Ventilation", "engineering101", "2000", "Open-Loop Ventilation"),
            ("kick-life-support-upgrade-PressurizedCabin", "Pressurized Cabins", "survivability", "0", "Pressurized Cabin"),
            ("kick-life-support-upgrade-LiOH", "LiOH Scrubber Systems", "survivability", "4000", "LiOH Scrubber"),
            ("kick-life-support-upgrade-Zeolite", "Zeolite Molecular Sieve Systems", "spaceExploration", "8000", "Zeolite Molecular Sieve"),
            ("kick-life-support-upgrade-SolidAmine", "Solid Amine Swingbed Systems", "advExploration", "16000", "Solid Amine Swingbed"),
            ("kick-life-support-upgrade-AirCooling", "Air Cooling Systems", "engineering101", "500", "Air Cooling"),
            ("kick-life-support-upgrade-WaterEvaporator", "Water Evaporator Systems", "engineering101", "1000", "Water Evaporator"),
            ("kick-life-support-upgrade-CoolantLoop", "Pumped Coolant Loop Systems", "survivability", "2000", "Pumped Coolant Loop"),
        ],
        [2800, 1900, 1500, 900, 2260],
    )

    doc.add_heading("How to patch your own upgrades", level=2)
    doc.add_paragraph(
        "1. Define a new PARTUPGRADE in 00_Upgrades.cfg or in your own patch file.\n"
        "2. Set techRequired to the target tech node and entryCost to the desired unlock cost.\n"
        "3. Reference that upgrade from the relevant B9 SUBTYPE with upgradeRequired = your-upgrade-name.\n"
        "4. If needed, add matching mass, cost, resource, or module DATA changes inside the same subtype patch."
    )
    add_note(
        doc,
        "Baseline subtypes without upgradeRequired remain available immediately. In the current setup, that means Unpressurized Cabin and Passive cooling are the always-on starting options."
    )


def main():
    doc = Document()
    style_document(doc)
    add_title(
        doc,
        "KICK Life Support Config and Patching Guide",
        "Reference for mod developers and players who want to edit values, write ModuleManager patches, or add new B9 subtype variants.",
    )
    doc.add_paragraph(
        "This guide describes the current local config structure in the KickLifeSupportFork workspace. "
        "It focuses on the fields and patch points you are most likely to edit."
    )

    add_settings_section(doc)
    add_b9_section(doc)
    add_tech_tree_section(doc)

    try:
        doc.save(OUTPUT)
        print(str(OUTPUT))
    except PermissionError:
        fallback = OUTPUT.with_name(
            f"{OUTPUT.stem}-{datetime.now().strftime('%Y%m%d-%H%M%S')}{OUTPUT.suffix}"
        )
        doc.save(fallback)
        print(str(fallback))


if __name__ == "__main__":
    main()
